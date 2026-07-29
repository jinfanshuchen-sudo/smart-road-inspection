from pathlib import Path
from docx import Document

SOURCE = Path(r"C:\Users\61700\Desktop\智慧道路巡检与路灯环境感知系统.docx")


def color_hex(run):
    color = run.font.color
    if color is None:
        return None
    if color.type is None:
        return None
    try:
        return str(color.rgb) if color.rgb else None
    except Exception:
        return None


def red_text(paragraph):
    pieces = []
    for run in paragraph.runs:
        color = color_hex(run)
        if color and color.upper() in {"FF0000", "FF0000"}:
            pieces.append(run.text)
    return "".join(pieces).strip()


def collect_paragraphs(parent, prefix, output):
    for index, paragraph in enumerate(parent.paragraphs, 1):
        red = red_text(paragraph)
        if red:
            output.append((f"{prefix} paragraph {index}", paragraph.text.strip(), red))
    for table_index, table in enumerate(parent.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            for cell_index, cell in enumerate(row.cells, 1):
                for paragraph_index, paragraph in enumerate(cell.paragraphs, 1):
                    red = red_text(paragraph)
                    if red:
                        label = (
                            f"{prefix} table {table_index}, row {row_index}, "
                            f"cell {cell_index}, paragraph {paragraph_index}"
                        )
                        output.append((label, paragraph.text.strip(), red))


document = Document(SOURCE)
entries = []
collect_paragraphs(document, "Body", entries)
for section_index, section in enumerate(document.sections, 1):
    collect_paragraphs(section.header, f"Section {section_index} header", entries)
    collect_paragraphs(section.footer, f"Section {section_index} footer", entries)

print(f"Red-text entries: {len(entries)}")
for label, full, red in entries:
    print("\n---", label)
    print("FULL:", full)
    print("RED :", red)

outline = []
for index, paragraph in enumerate(document.paragraphs, 1):
    text = paragraph.text.strip()
    if text:
        outline.append(f"P{index} [{paragraph.style.name}] {text}")
for table_index, table in enumerate(document.tables, 1):
    outline.append(f"\n[TABLE {table_index}]")
    for row_index, row in enumerate(table.rows, 1):
        cells = [" | ".join(p.text.strip() for p in cell.paragraphs if p.text.strip()) for cell in row.cells]
        outline.append(f"R{row_index}: " + " || ".join(cells))
Path("document_outline_utf8.txt").write_text("\n".join(outline), encoding="utf-8")
