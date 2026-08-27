from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent.parent / "客户交付包_智慧道路巡检系统" / "智慧道路巡检系统_甲方使用说明书.docx"

BLUE = "2E74B5"
DARK = "1F4D78"
NAVY = "0B2545"
MUTED = "5E6B78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GOLD = "7A5A00"
RED = "9B1C1C"


def set_font(run, name="Microsoft YaHei", size=None, color=None, bold=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    tr_pr.append(hdr)


def table_borders(table, color="D6DEE8", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        item = borders.find(tag)
        if item is None:
            item = OxmlElement(f"w:{edge}")
            borders.append(item)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), size)
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), color)


def add_text(p, text, size=11, color=NAVY, bold=False, align=None, before=0, after=6, line=1.25):
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text, level=1):
    style = doc.styles[f"Heading {level}"]
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_with_next = True
    add_text(
        p,
        text,
        size={1: 16, 2: 13, 3: 12}[level],
        color={1: BLUE, 2: BLUE, 3: DARK}[level],
        bold=True,
        before={1: 16, 2: 12, 3: 8}[level],
        after={1: 8, 2: 6, 3: 4}[level],
    )
    return p


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        add_text(p, item, size=11, color=NAVY, after=4, line=1.25)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        add_text(p, item, size=11, color=NAVY, after=4, line=1.25)


def add_label_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [2700, 6660])
    table_borders(table)
    for label, detail in rows:
        cells = table.add_row().cells
        shade(cells[0], LIGHT_BLUE)
        for cell in cells:
            cell.text = ""
        add_text(cells[0].paragraphs[0], label, size=10.5, color=DARK, bold=True, after=0, line=1.1)
        add_text(cells[1].paragraphs[0], detail, size=10.5, color=NAVY, after=0, line=1.1)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_status_table(doc):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2480, 2500, 4380])
    table_borders(table)
    headers = ["页面显示", "代表含义", "处理方式"]
    for cell, text in zip(table.rows[0].cells, headers):
        shade(cell, LIGHT_BLUE)
        cell.text = ""
        add_text(cell.paragraphs[0], text, size=10.5, color=DARK, bold=True, after=0, line=1.1)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("已连接（绿色）", "网页已连接树莓派本地服务", "正常使用。"),
        ("待接入 / 未启用", "对应设备暂未上电、未连接或暂未上报数据", "检查该设备是否已供电并连接 Hula-Battle。"),
        ("网页无法打开", "电脑或手机未连到正确 WiFi，或树莓派未启动", "确认 Hula-Battle、树莓派供电；等待 1 分钟后重试。"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values):
            cell.text = ""
            add_text(cell.paragraphs[0], text, size=10.2, color=NAVY, after=0, line=1.1)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, text, fill="F4F6F9", title_color=DARK):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    table_borders(table, color="CFD8E3", size="6")
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    add_text(p, title, size=10.5, color=title_color, bold=True, after=2, line=1.1)
    p2 = cell.add_paragraph()
    add_text(p2, text, size=10.5, color=NAVY, after=0, line=1.15)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run2 = paragraph.add_run(" 页")
    set_font(run2, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color in ((1, 16, BLUE), (2, 13, BLUE), (3, 12, DARK)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    hp = section.header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(hp, "智慧道路巡检系统  |  甲方使用说明", size=9, color=MUTED, after=0, line=1)
    fp = section.footer.paragraphs[0]
    fp.text = ""
    add_page_number(fp)


def build():
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    add_text(p, "客户交付使用说明", size=10.5, color=GOLD, bold=True, after=4)
    p = doc.add_paragraph()
    add_text(p, "智慧道路巡检系统", size=27, color=NAVY, bold=True, after=4)
    p = doc.add_paragraph()
    add_text(p, "本地离线监测与巡检平台", size=14, color=MUTED, after=18)

    add_callout(
        doc,
        "最简单的使用方式",
        "给路由器、树莓派及现场设备上电；手机或电脑连接 Hula-Battle WiFi；在浏览器打开 http://192.168.31.66:5055。无需安装软件，无需连接互联网。",
        fill="EAF2F8",
        title_color=DARK,
    )

    add_heading(doc, "一、系统用途", 1)
    p = doc.add_paragraph()
    add_text(
        p,
        "本系统用于现场展示和查看气象、积水、人车接近、灯塔状态及无人机巡检相关信息。树莓派作为现场主机，持续运行监测网页和本地数据服务；管理人员可通过同一 WiFi 内的任意手机、平板或电脑访问网页。",
    )

    add_heading(doc, "二、使用前请确认", 1)
    add_label_table(
        doc,
        [
            ("现场 WiFi", "Hula-Battle"),
            ("访问地址", "http://192.168.31.66:5055"),
            ("是否需要互联网", "不需要。本系统在现场局域网内运行。"),
            ("树莓派状态", "树莓派需保持供电，并放置在 Hula-Battle WiFi 覆盖范围内。"),
            ("现场设备", "ESP32、气象站、积水模块、灯带及无人机按原有连接方式接好并上电。"),
        ],
    )

    add_heading(doc, "三、日常启动与进入系统", 1)
    add_numbered(
        doc,
        [
            "确认 Hula-Battle 路由器已上电并正常工作。",
            "给树莓派供电。树莓派会自动连接 Hula-Battle，并自动启动系统；请等待约 1 分钟。",
            "按现场需要给 ESP32、气象站、积水模块、灯带等设备上电。",
            "使用任意手机、平板或电脑连接 Hula-Battle WiFi。",
            "打开浏览器，在地址栏输入 http://192.168.31.66:5055。",
        ],
    )
    add_callout(
        doc,
        "重要提示",
        "访问设备必须连接 Hula-Battle。不要使用手机流量、其他办公室 WiFi 或 VPN 访问本系统。",
        fill="FFF8E8",
        title_color=GOLD,
    )

    add_heading(doc, "四、网页状态怎么看", 1)
    add_status_table(doc)

    add_heading(doc, "五、正常使用说明", 1)
    add_heading(doc, "1. 监测数据", 2)
    add_bullets(
        doc,
        [
            "气象站区域显示风速、风向、雨量、光照、温度、湿度、气压等数据。",
            "人车接近区域显示检测启用状态、在线状态、最近距离和报警等级。",
            "积水检查区域显示积水监测状态和相关风险信息。",
            "综合告警区域用于汇总当前现场告警与最近更新时间。",
        ],
    )
    add_heading(doc, "2. 无人机巡检", 2)
    p = doc.add_paragraph()
    add_text(
        p,
        "无人机操作前，请先确认无人机、树莓派和操作终端均连接 Hula-Battle，并在安全、空旷且符合现场管理要求的区域内操作。无人机起飞、巡航、返航和降落应由具备相应操作资格的人员执行。",
    )

    add_heading(doc, "六、常见问题处理", 1)
    add_heading(doc, "网页打不开", 2)
    add_numbered(
        doc,
        [
            "确认访问设备已连接 Hula-Battle。",
            "确认树莓派电源已接通，等待 1 分钟后刷新网页。",
            "确认地址输入完整且正确：http://192.168.31.66:5055。",
            "若仍无法访问，请联系项目交付方处理；请勿自行修改树莓派 IP、路由器 WiFi 名称或密码。",
        ],
    )
    add_heading(doc, "网页能打开但没有数据", 2)
    add_bullets(
        doc,
        [
            "检查 ESP32、气象站、积水模块、灯带等是否已供电。",
            "检查 ESP32 是否已连接 Hula-Battle。",
            "确认各设备接线保持交付时的原有连接方式，不要将气象站 12V 电源接入树莓派、ESP32、LD06 或灯带。",
            "设备刚上电时可等待数十秒，待数据开始上报后页面会自动更新。",
        ],
    )
    add_heading(doc, "需要长期停用现场设备", 2)
    p = doc.add_paragraph()
    add_text(
        p,
        "请先停止无人机和现场用电设备。树莓派不建议频繁直接拔电；如需关闭树莓派，请由项目维护人员执行安全关机后再断电。",
    )

    add_heading(doc, "七、交付后的注意事项", 1)
    add_bullets(
        doc,
        [
            "树莓派是本系统的现场主机，日常使用无需连接原开发笔记本。",
            "不要更改 Hula-Battle 的 WiFi 名称、密码或网段；如确需更改，请联系项目交付方重新配置。",
            "不要修改树莓派固定地址 192.168.31.66。",
            "本系统为本地局域网系统，现场断网不影响本地网页和设备之间的基本通信。",
        ],
    )
    add_callout(
        doc,
        "交付验收建议",
        "树莓派上电后，使用另一台手机或电脑连接 Hula-Battle 并成功打开网页；再逐项确认气象、积水、人车接近、灯带及无人机相关功能是否按现场配置显示或响应。",
        fill="EAF2F8",
        title_color=DARK,
    )

    doc.core_properties.title = "智慧道路巡检系统甲方使用说明书"
    doc.core_properties.subject = "本地离线监测与巡检平台交付说明"
    doc.core_properties.author = "项目交付方"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
