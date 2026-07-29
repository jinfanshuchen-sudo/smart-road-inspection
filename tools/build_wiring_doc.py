from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\Users\61700\Desktop\pyhulax-main")
OUT = Path(r"C:\Users\61700\Desktop\智慧道路巡检系统_硬件接线说明与接线图.docx")
ASSET_DIR = ROOT / "docx_diagrams"
ASSET_DIR.mkdir(exist_ok=True)
DIAGRAM = ASSET_DIR / "hardware_wiring_diagram.png"

NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "DCE6F1"
PALE_BLUE = "EEF4FA"
GREY = "F2F4F7"
MID_GREY = "D9E2F3"
RED = "C00000"
YELLOW = "F2C94C"
GREEN = "2E8B57"
BLACK = "202020"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def add_text(cell, text, bold=False, color=BLACK, size=8.7, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc, headers, rows, widths, font_size=8.7):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        add_text(table.rows[0].cells[i], header, True, WHITE, font_size, WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if r_idx % 2 == 0:
                set_cell_shading(cells[i], PALE_BLUE)
            add_text(cells[i], value, False, BLACK, font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def set_paragraph(p, before=0, after=5, line=1.25):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_para(doc, text="", bold=False, color=BLACK, size=9.5, before=0, after=5, indent=0):
    p = doc.add_paragraph()
    set_paragraph(p, before, after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22 + 0.18 * level)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    set_paragraph(p, 0, 3, 1.2)
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(9.3)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph(p, 10 if level == 1 else 6, 4, 1.1)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(14 if level == 1 else 11.2)
    r.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else BLUE)
    return p


def add_note(doc, label, text, fill="FFF2CC"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [7920])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph(p, 0, 0, 1.2)
    r1 = p.add_run(label + "：")
    r1.bold = True
    r1.font.name = "Microsoft YaHei"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r1.font.size = Pt(9.2)
    r1.font.color.rgb = RGBColor.from_string(RED)
    r2 = p.add_run(text)
    r2.font.name = "Microsoft YaHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def font(size, bold=False):
    for name in [r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\simhei.ttf"]:
        if Path(name).exists():
            return ImageFont.truetype(name, size)
    return ImageFont.load_default()


def image_color(color):
    return color if color.startswith("#") else f"#{color}"


def draw_box(draw, box, title, lines, fill, outline=NAVY):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=22, fill=image_color(fill), outline=image_color(outline), width=4)
    title_font = font(34, True)
    body_font = font(25)
    draw.text((x1 + 24, y1 + 18), title, font=title_font, fill=image_color(NAVY))
    yy = y1 + 68
    for line in lines:
        draw.text((x1 + 24, yy), line, font=body_font, fill=image_color(BLACK))
        yy += 34


def draw_arrow(draw, start, end, color, label, vertical_offset=0):
    color = image_color(color)
    draw.line([start, end], fill=color, width=8)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - 22 * direction, ey - 13), (ex - 22 * direction, ey + 13)]
    else:
        direction = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 13, ey - 22 * direction), (ex + 13, ey - 22 * direction)]
    draw.polygon(pts, fill=color)
    mx, my = (sx + ex) / 2, (sy + ey) / 2
    f = font(23, True)
    bbox = draw.textbbox((0, 0), label, font=f)
    pad = 7
    draw.rounded_rectangle((mx - (bbox[2]-bbox[0])/2-pad, my-14+vertical_offset, mx + (bbox[2]-bbox[0])/2+pad, my + 18+vertical_offset), radius=6, fill="#FFFFFF")
    draw.text((mx-(bbox[2]-bbox[0])/2, my-10+vertical_offset), label, font=f, fill=color)


def create_diagram():
    img = Image.new("RGB", (2400, 1500), "#FFFFFF")
    d = ImageDraw.Draw(img)
    title_font = font(50, True)
    d.text((70, 45), "智慧道路巡检系统 - 硬件总接线图", font=title_font, fill=image_color(NAVY))
    d.text((72, 112), "颜色：黄色=5V电源  红色=12V电源  黑色=公共地  蓝色=TTL信号  绿色=RS485 A/B  灰色=USB数据", font=font(24), fill=image_color(BLACK))

    draw_box(d, (70, 260, 520, 520), "外接口 A：电脑", ["USB 数据线", "数据/调试", "可为 ESP32 供电"], "EEF4FA")
    draw_box(d, (840, 245, 1440, 610), "ESP32-S3 主控", ["GPIO48 -> 灯带 DIN", "GPIO18 <- LD06 TX", "GPIO4 <- 水位 AO", "GPIO5 <- 水位 DO", "GPIO15 -> RS485 RXD", "GPIO16 <- RS485 TXD"], "DCE6F1")
    draw_box(d, (1780, 210, 2320, 470), "LD06 雷达", ["TX -> GPIO18", "VCC -> 外接 5V", "GND -> 公共地"], "F2F4F7")
    draw_box(d, (1780, 570, 2320, 830), "WS2812 灯带", ["DIN <- GPIO48", "5V -> 外接 5V", "GND -> 公共地"], "F2F4F7")
    draw_box(d, (1780, 930, 2320, 1210), "水位检测模块", ["VCC -> ESP32 3V3", "AO -> GPIO4", "DO -> GPIO5", "GND -> 公共地"], "F2F4F7")
    draw_box(d, (70, 900, 560, 1180), "外接口 B：5V 电源", ["+5V -> LD06 VCC", "+5V -> 灯带 5V", "0V -> 公共地"], "FFF2CC", outline="B58105")
    draw_box(d, (800, 935, 1450, 1235), "TTL-RS485 转换模块", ["RXD <- GPIO15", "TXD -> GPIO16", "A/B -> 气象站 A/B", "VCC 需兼容 3.3V"], "E2F0D9", outline=GREEN)
    draw_box(d, (1740, 1280, 2320, 1460), "气象站 RS485", ["A/B: 通信端", "+12V/0V: 独立供电"], "E2F0D9", outline=GREEN)
    draw_box(d, (70, 1260, 560, 1460), "外接口 C：12V 电源", ["+12V -> 气象站 V+", "0V -> 气象站 V-", "不接 ESP32 5V"], "FCE4D6", outline=RED)

    draw_arrow(d, (520, 385), (840, 385), "6B7280", "USB 数据")
    draw_arrow(d, (1440, 350), (1780, 350), BLUE, "GPIO18 / LD06 TX")
    draw_arrow(d, (1440, 515), (1780, 690), BLUE, "GPIO48 / DIN")
    draw_arrow(d, (1440, 570), (1780, 1050), BLUE, "GPIO4 / GPIO5")
    draw_arrow(d, (560, 1015), (1780, 400), "B58105", "+5V")
    draw_arrow(d, (560, 1100), (1780, 750), "B58105", "+5V")
    draw_arrow(d, (1440, 1080), (1740, 1360), GREEN, "RS485 A/B")
    draw_arrow(d, (560, 1360), (1740, 1405), RED, "+12V / 0V")
    d.line([(650, 1260), (650, 1340), (760, 1340), (760, 610), (1780, 610)], fill="#111111", width=6)
    d.text((570, 1380), "公共地：ESP32 GND、LD06 0V、灯带 0V、水位 GND、TTL-RS485 GND", font=font(21, True), fill="#111111")
    d.text((80, 180), "注：RS485 方向控制请使用自动收发模块；若是普通 MAX485，DE/RE 需由 GPIO7 控制，现有程序未驱动该脚。", font=font(22), fill=image_color(RED))
    img.save(DIAGRAM)


def draw_route(draw, points, color, width=7):
    """Draw one routed line with an arrow only at its final segment."""
    color = image_color(color)
    draw.line(points, fill=color, width=width, joint="curve")
    sx, sy = points[-2]
    ex, ey = points[-1]
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        arrow = [(ex, ey), (ex - 20 * direction, ey - 12), (ex - 20 * direction, ey + 12)]
    else:
        direction = 1 if ey > sy else -1
        arrow = [(ex, ey), (ex - 12, ey - 20 * direction), (ex + 12, ey - 20 * direction)]
    draw.polygon(arrow, fill=color)


def draw_line_label(draw, xy, text, color, size=20):
    """Put labels beside a wire, never directly on top of it."""
    f = font(size, True)
    color = image_color(color)
    bbox = draw.textbbox((0, 0), text, font=f)
    w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
    x, y = xy
    draw.rounded_rectangle((x - 7, y - 5, x + w + 7, y + h + 5), radius=6, fill="#FFFFFF", outline="#D1D5DB", width=1)
    draw.text((x, y), text, font=f, fill=color)


def create_clean_diagram():
    """Create a compact, non-overlapping wiring map for the Word document."""
    img = Image.new("RGB", (2400, 1330), "#FFFFFF")
    d = ImageDraw.Draw(img)

    d.text((70, 35), "\u667a\u6167\u9053\u8def\u5de1\u68c0\u7cfb\u7edf \u786c\u4ef6\u603b\u63a5\u7ebf\u56fe", font=font(48, True), fill=image_color(NAVY))
    d.text((72, 100), "\u9ec4\u8272=5V\u7535\u6e90   \u7ea2\u8272=12V\u7535\u6e90   \u9ed1\u8272=\u516c\u5171\u5730   \u84dd\u8272=TTL\u4fe1\u53f7   \u7eff\u8272=RS485 A/B   \u7070\u8272=USB\u6570\u636e", font=font(23), fill=image_color(BLACK))
    d.text((72, 150), "\u6ce8\uff1aRS485\u5efa\u8bae\u4f7f\u7528\u81ea\u52a8\u6536\u53d1\u8f6c\u6362\u6a21\u5757\uff1b\u666e\u901a MAX485 \u9700 DE/RE \u65b9\u5411\u63a7\u5236\u3002", font=font(21), fill=image_color(RED))

    draw_box(d, (70, 235, 500, 465), "\u5916\u63a5\u53e3 A\uff1a\u7535\u8111", ["USB \u6570\u636e\u7ebf", "\u6570\u636e/\u8c03\u8bd5", "\u53ef\u4e3a ESP32 \u4f9b\u7535"], "EEF4FA")
    draw_box(d, (820, 215, 1450, 595), "ESP32-S3 \u4e3b\u63a7", ["GPIO48 -> \u706f\u5e26 DIN", "GPIO18 <- LD06 TX", "GPIO4 <- \u6c34\u4f4d AO", "GPIO5 <- \u6c34\u4f4d DO", "GPIO15 -> RS485 RXD", "GPIO16 <- RS485 TXD"], "DCE6F1")
    draw_box(d, (1800, 205, 2320, 425), "LD06 \u96f7\u8fbe", ["TX -> GPIO18", "VCC -> \u5916\u63a5 5V", "GND -> \u516c\u5171\u5730"], "F2F4F7")
    draw_box(d, (1800, 510, 2320, 730), "WS2812 \u706f\u5e26", ["DIN <- GPIO48", "5V -> \u5916\u63a5 5V", "GND -> \u516c\u5171\u5730"], "F2F4F7")
    draw_box(d, (1800, 815, 2320, 1050), "\u6c34\u4f4d\u68c0\u6d4b\u6a21\u5757", ["VCC -> ESP32 3V3", "AO -> GPIO4", "DO -> GPIO5", "GND -> \u516c\u5171\u5730"], "F2F4F7")
    draw_box(d, (70, 720, 500, 950), "\u5916\u63a5\u53e3 B\uff1a5V \u7535\u6e90", ["+5V -> LD06 VCC", "+5V -> \u706f\u5e26 5V", "0V -> \u516c\u5171\u5730"], "FFF2CC", outline="B58105")
    draw_box(d, (800, 755, 1480, 1015), "TTL-RS485 \u8f6c\u6362\u6a21\u5757", ["RXD <- GPIO15", "TXD -> GPIO16", "A/B -> \u6c14\u8c61\u7ad9 A/B", "VCC \u9700\u517c\u5bb9 3.3V"], "E2F0D9", outline=GREEN)
    draw_box(d, (70, 1085, 500, 1280), "\u5916\u63a5\u53e3 C\uff1a12V \u7535\u6e90", ["+12V -> \u6c14\u8c61\u7ad9 V+", "0V -> \u6c14\u8c61\u7ad9 V-", "\u4e0d\u63a5 ESP32 5V"], "FCE4D6", outline=RED)
    draw_box(d, (1770, 1100, 2320, 1280), "\u6c14\u8c61\u7ad9 RS485", ["A/B\uff1a\u901a\u4fe1\u7aef", "+12V/0V\uff1a\u72ec\u7acb\u4f9b\u7535"], "E2F0D9", outline=GREEN)

    # USB and TTL signals: each path uses an unused corridor.
    draw_route(d, [(500, 350), (820, 350)], "6B7280")
    draw_line_label(d, (605, 315), "USB \u6570\u636e", "6B7280")
    draw_route(d, [(1450, 330), (1800, 330)], BLUE)
    draw_line_label(d, (1510, 288), "GPIO18 / LD06 TX", BLUE)
    draw_route(d, [(1450, 450), (1660, 450), (1660, 620), (1800, 620)], BLUE)
    draw_line_label(d, (1495, 470), "GPIO48 / DIN", BLUE)
    draw_route(d, [(1450, 550), (1600, 550), (1600, 920), (1800, 920)], BLUE)
    draw_line_label(d, (1475, 720), "GPIO4 / GPIO5", BLUE)

    # External 5V gets its own upper and middle routes, away from signal wires.
    draw_route(d, [(500, 785), (620, 785), (620, 185), (1700, 185), (1700, 280), (1800, 280)], "B58105")
    draw_line_label(d, (650, 205), "+5V -> LD06", "B58105")
    draw_route(d, [(500, 875), (650, 875), (650, 705), (1700, 705), (1700, 675), (1800, 675)], "B58105")
    draw_line_label(d, (960, 665), "+5V -> \u706f\u5e26", "B58105")

    # RS485 and 12V are separated vertically, so their labels remain readable.
    draw_route(d, [(1120, 595), (1120, 680), (1240, 680), (1240, 755)], GREEN)
    draw_line_label(d, (1025, 645), "UART \u4e32\u53e3", GREEN)
    draw_route(d, [(1480, 875), (1660, 875), (1660, 1190), (1770, 1190)], GREEN)
    draw_line_label(d, (1500, 920), "RS485 A/B", GREEN)
    draw_route(d, [(500, 1175), (1570, 1175), (1570, 1245), (1770, 1245)], RED)
    draw_line_label(d, (780, 1125), "+12V / 0V \u2192 \u6c14\u8c61\u7ad9", RED)

    # Dedicated ground bus sits below all signal routes and never shares their labels.
    ground_y = 1065
    d.line([(600, ground_y), (1710, ground_y)], fill="#111111", width=6)
    for points in [
        [(500, 920), (600, 920), (600, ground_y)],
        [(1040, 595), (1040, ground_y)],
        [(1240, 1015), (1240, ground_y)],
        [(1800, 390), (1730, 390), (1730, ground_y)],
        [(1800, 720), (1700, 720), (1700, ground_y)],
        [(1800, 1015), (1710, 1015), (1710, ground_y)],
    ]:
        d.line(points, fill="#111111", width=5)
    draw_line_label(d, (755, 1088), "\u7cfb\u7edf\u516c\u5171\u5730\uff1aESP32 GND / LD06 0V / \u706f\u5e26 0V / \u6c34\u4f4d GND / TTL-RS485 GND", "111111", 17)
    img.save(DIAGRAM)


def build_doc():
    create_clean_diagram()
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.27)
    sec.page_height = Inches(11.69)
    sec.top_margin = Inches(0.58)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.62)
    sec.right_margin = Inches(0.62)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(9.5)

    # Header/footer
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("智慧道路巡检与路灯环境感知系统  |  硬件接线说明")
    hr.font.name = "Microsoft YaHei"
    hr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor.from_string("6B7280")
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("\u786c\u4ef6\u63a5\u7ebf\u6838\u5bf9\u7248")
    fr.font.name = "Microsoft YaHei"
    fr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string("6B7280")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(title, 6, 2, 1.0)
    r = title.add_run("智慧道路巡检系统")
    r.bold = True; r.font.name = "Microsoft YaHei"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei"); r.font.size = Pt(24); r.font.color.rgb = RGBColor.from_string(NAVY)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(sub, 0, 13, 1.0)
    r = sub.add_run("\u786c\u4ef6\u63a5\u7ebf\u8bf4\u660e\u4e0e\u603b\u63a5\u7ebf\u56fe")
    r.font.name = "Microsoft YaHei"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei"); r.font.size = Pt(12); r.font.color.rgb = RGBColor.from_string(BLUE)

    add_note(doc, "适用范围", "本手册覆盖 ESP32-S3、LD06 雷达、水位检测模块、WS2812 灯带、RS485 气象站、TTL-RS485 转换模块及电脑 USB 数据连接的供电、信号与共地关系。", "EAF2F8")
    add_heading(doc, "1. 三路外接接口与供电原则")
    add_table(doc, ["外接口", "接到哪里", "用途", "关键要求"], [
        ["A. ESP32 USB 数据线", "电脑 USB <-> ESP32-S3 USB 口", "烧录程序、串口调试、电脑与 ESP32 数据连接", "使用可传数据的 USB 线；只用电脑 USB 时，不要让 LD06/灯带从该口取大电流。"],
        ["B. LD06/灯带 5V 电源", "5V+ 分两路至 LD06 VCC、灯带 5V；5V- 接公共地", "给 LD06 和灯带提供独立 5V", "5V- 必须与 ESP32 GND 共地；不得把 5V 接进 ESP32 的 3V3 引脚。"],
        ["C. 气象站 12V 电源", "12V+ / 12V- 只接气象站电源端", "给 RS485 气象站供电", "12V 严禁接入 ESP32、LD06、灯带或 TTL-RS485 的 VCC。"],
    ], [1320, 2100, 1950, 2550], 8.4)
    add_para(doc, "公共地规则：ESP32 GND、LD06 的 5V-、灯带 5V-、水位模块 GND、TTL-RS485 模块 GND 必须连接到同一公共地。气象站的 12V- 仅按气象站/RS485模块说明处理，不要未经确认直接接到 ESP32 5V 线路。", True, RED, 9.2, 4, 7)

    add_heading(doc, "2. 彩色总接线图")
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.space_before = Pt(2)
    pic.paragraph_format.space_after = Pt(0)
    pic.paragraph_format.keep_with_next = True
    run = pic.add_run()
    run.add_picture(str(DIAGRAM), width=Inches(6.95))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(caption, 0, 3, 1.0)
    c = caption.add_run("\u56fe 1  \u7cfb\u7edf\u786c\u4ef6\u603b\u4f53\u63a5\u7ebf\u5173\u7cfb\u56fe")
    c.italic = True; c.font.name = "Microsoft YaHei"; c._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei"); c.font.size = Pt(8.5); c.font.color.rgb = RGBColor.from_string("6B7280")
    add_heading(doc, "3. ESP32-S3 引脚总表")
    add_para(doc, "以下引脚取自当前项目固件的 pinmap.h；文档按现有程序工作方式整理。", False, "4B5563", 9.3, 0, 6)
    add_table(doc, ["ESP32-S3 引脚", "连接模块/线", "方向", "说明"], [
        ["GPIO48", "WS2812 灯带 DIN", "ESP32 -> 灯带", "当前状态灯/灯带数据输出。灯带必须另供 5V，并与 ESP32 共地。"],
        ["GPIO18", "LD06 TX", "LD06 -> ESP32", "LD06 串口数据输入（RX）。只接这一根数据线；当前程序不向 LD06 发串口数据。"],
        ["GPIO4", "水位模块 AO", "水位模块 -> ESP32", "水位模拟量输入。传感器 AO 输出不得超过 3.3V；若模块以 5V 供电，需先确认 AO 电压或加分压。"],
        ["GPIO5", "水位模块 DO", "水位模块 -> ESP32", "水位数字阈值输出。与 AO 一起用于水位采样/状态判断。"],
        ["GPIO15", "TTL-RS485 RXD", "ESP32 -> RS485", "ESP32 串口发送线（UART TX）。"],
        ["GPIO16", "TTL-RS485 TXD", "RS485 -> ESP32", "ESP32 串口接收线（UART RX）。"],
        ["GPIO7", "RS485 DE/RE（仅普通 MAX485 时）", "ESP32 -> RS485", "当前程序未驱动 GPIO7；优先使用自动收发 TTL-RS485 模块。"],
        ["GPIO17 / 13 / 14 / 8 / 9", "保留", "-", "当前版本不作为 LD06、灯带、水位或气象站的实际接线端。不要擅自接入。"],
    ], [1460, 2050, 1600, 2810], 8.3)
    add_note(doc, "重要", "ESP32 IO 逻辑电平为 3.3V。任何可能输出 5V 的 AO、DO、TXD 信号，先核实电平；未确认前不得直接接 GPIO。", "FCE4D6")

    add_heading(doc, "4. LD06 雷达、灯带和水位模块")
    add_heading(doc, "4.1 LD06 雷达（外接 5V 供电）", 2)
    add_table(doc, ["LD06 端", "接到哪里", "接线颜色建议", "说明"], [
        ["VCC / 5V", "外接口 B 的 +5V", "黄/红", "LD06 使用独立稳定 5V。"],
        ["GND", "外接口 B 的 0V，并与 ESP32 GND 相连", "黑", "必须与 ESP32 共地，否则 UART 数据无法可靠识别。"],
        ["TX", "ESP32 GPIO18", "蓝", "LD06 数据输出 -> ESP32 接收。"],
        ["RX", "不接", "-", "当前程序只接收 LD06 数据，不需要 ESP32 -> LD06 的发送线。"],
    ], [1220, 2820, 1320, 2560], 8.4)
    add_heading(doc, "4.2 WS2812 / RGB 灯带（取代 ESP32 板载大灯）", 2)
    add_table(doc, ["灯带端", "接到哪里", "说明"], [
        ["5V", "外接口 B 的 +5V", "灯带不可依赖 ESP32 板载 5V 供电。按灯带数量预留足够电流。"],
        ["GND", "外接口 B 的 0V，并与 ESP32 GND 共地", "共地是灯带数据稳定显示的必要条件。"],
        ["DIN", "ESP32 GPIO48", "只接灯带输入端 DIN；若灯带还有 DOUT，不接回 ESP32。"],
    ], [1220, 2700, 4000], 8.4)
    add_heading(doc, "4.3 水位检测模块", 2)
    add_table(doc, ["水位模块端", "接到哪里", "说明"], [
        ["VCC", "ESP32 3V3", "当前工程建议优先使用 3.3V 供电，确保 AO/DO 对 GPIO 安全。"],
        ["GND", "ESP32 GND", "与系统公共地连接。"],
        ["AO", "ESP32 GPIO4", "模拟水位值。"],
        ["DO", "ESP32 GPIO5", "数字阈值状态。"],
    ], [1220, 2700, 4000], 8.4)

    add_heading(doc, "5. RS485 气象站接线")
    add_para(doc, "气象站由独立 12V 电源供电，数据经 TTL-RS485 转换模块进入 ESP32。12V 只到气象站电源端，不能进入 TTL-RS485 模块 VCC。", True, RED, 9.2, 0, 6)
    add_table(doc, ["部件端子", "接到哪里", "线路性质", "说明"], [
        ["气象站 V+", "外接口 C 的 +12V", "电源", "仅给气象站供电。"],
        ["气象站 V-", "外接口 C 的 0V", "电源", "仅给气象站供电。"],
        ["气象站 A", "TTL-RS485 模块 A", "RS485 差分通信", "A 对 A；若无数据再在断电状态下核对 A/B 标识。"],
        ["气象站 B", "TTL-RS485 模块 B", "RS485 差分通信", "B 对 B。"],
        ["ESP32 GPIO15", "TTL-RS485 模块 RXD", "TTL 串口", "ESP32 TX -> 模块 RXD。"],
        ["ESP32 GPIO16", "TTL-RS485 模块 TXD", "TTL 串口", "模块 TXD -> ESP32 RX。"],
        ["ESP32 3V3", "TTL-RS485 模块 VCC", "模块逻辑供电", "仅适用于支持 3.3V 逻辑的 TTL-RS485 模块。"],
        ["ESP32 GND", "TTL-RS485 模块 GND", "公共地", "用于 TTL 串口参考地。"],
    ], [1660, 2400, 1600, 2260], 8.25)
    add_note(doc, "RS485 模块选择", "推荐自动收发型 TTL-RS485 模块（无需 DE/RE 控制）。若使用普通 MAX485 类型，DE 与 RE 通常需并接 GPIO7，但当前固件未主动切换该脚，必须先改程序后再接线。", "FFF2CC")

    add_heading(doc, "6. 实物接线顺序与上电检查")
    numbered = [
        "断开所有电源：电脑 USB、LD06 5V、气象站 12V 都先拔掉。",
        "先完成低压信号线：LD06 TX -> GPIO18；水位 AO/DO -> GPIO4/GPIO5；灯带 DIN -> GPIO48；TTL-RS485 RXD/TXD -> GPIO15/GPIO16。",
        "完成公共地：ESP32 GND、LD06 0V、灯带 0V、水位 GND、TTL-RS485 GND 相连。",
        "接外接口 B：外部 5V 分别接 LD06 VCC 和灯带 5V；确认极性无误。",
        "接外接口 C：外部 12V 只接气象站 V+/V-；再接气象站 A/B 到 TTL-RS485 模块 A/B。",
        "最后接外接口 A：电脑 USB 接 ESP32。确认没有短路、没有把 12V 接到 ESP32/灯带/LD06。",
        "上电后依次检查：ESP32 串口启动正常 -> 灯带状态变化 -> LD06 有距离数据 -> 水位 ADC 有变化 -> 气象站有 RS485 数据。",
    ]
    for i, item in enumerate(numbered, 1):
        p = doc.add_paragraph(style="List Number")
        set_paragraph(p, 0, 3, 1.2)
        r = p.add_run(item)
        r.font.name = "Microsoft YaHei"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei"); r.font.size = Pt(9.25)
    add_note(doc, "故障快速定位", "LD06 无数据先检查 5V 与共地，再检查 TX 是否接 GPIO18；灯带异常先检查 5V、电流能力、GPIO48 与共地；水位没有变化先核实 VCC=3V3、AO/DO 引脚；气象站无数据先检查 12V、A/B、TTL-RS485 是否为自动收发型。", "EAF2F8")

    add_heading(doc, "7. 接线核对清单")
    add_table(doc, ["核对项", "正确状态", "现场勾选"], [
        ["外部 5V", "仅供 LD06 与灯带，负极已与 ESP32 GND 共地", "□"],
        ["外部 12V", "仅接气象站 V+/V-，未接入 ESP32/LD06/灯带", "□"],
        ["LD06", "TX -> GPIO18；VCC=5V；GND 已共地", "□"],
        ["灯带", "DIN -> GPIO48；5V 独立供电；GND 已共地", "□"],
        ["水位模块", "VCC=3V3；AO -> GPIO4；DO -> GPIO5；GND 正确", "□"],
        ["RS485 气象站", "GPIO15 -> RXD；GPIO16 <- TXD；A/B 对应；模块支持 3.3V/自动收发", "□"],
        ["电脑 USB", "使用数据线连接 ESP32，电脑可识别串口", "□"],
    ], [1900, 4850, 1170], 8.35)

    add_heading(doc, "8. 现场端口标识与演示前检查")
    add_table(doc, ["端口", "对应线缆", "连接对象", "使用要点"], [
        ["A. USB口", "USB 数据线", "电脑 / ESP32-S3", "用于供电、烧录与串口调试；演示前确认电脑可识别串口。"],
        ["B. 5V口", "5V 两芯线", "LD06 与 WS2812 灯带", "+5V 只接 VCC/5V，0V 必须与 ESP32 GND 共地。"],
        ["C. 12V口", "12V 两芯线", "RS485 气象站", "只接气象站 V+/V-，不可连接到 ESP32、LD06 或灯带。"],
    ], [1260, 1540, 1800, 3320], 8.5)
    add_note(doc, "演示前 30 秒检查", "请在线缆上分别贴上 A-USB 数据、B-5V 供电、C-12V 气象站标签。上电后先确认三个端口极性、共地和插头牢固，再打开系统进行数据演示。", "EAF2F8")

    doc.core_properties.title = "智慧道路巡检系统硬件接线说明与接线图"
    doc.core_properties.subject = "ESP32-S3、LD06、灯带、水位、RS485气象站接线"
    doc.core_properties.author = "项目组"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
