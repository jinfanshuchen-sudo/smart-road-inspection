from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\61700\Desktop\�ǻ۵�·Ѳ����·�ƻ�����֪ϵͳ.docx")
OUTPUT = Path(r"C:\Users\61700\Desktop\pyhulax-main\�ǻ۵�·Ѳ����·�ƻ�����֪ϵͳ_�Ѳ�����ɫ�˶԰�.docx")

BLUE = RGBColor(0x00, 0x00, 0xFF)
LIGHT_BLUE_FILL = "DDEBF7"


def set_run_blue(run, bold=False, size=10.5):
    run.font.color.rgb = BLUE
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "����"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:eastAsia"), "����")


def insert_paragraph_after(paragraph):
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    return Paragraph(element, paragraph._parent)


def add_blue_paragraph_after(paragraph, text, *, bold=False, indent=False, space_after=3):
    new_paragraph = insert_paragraph_after(paragraph)
    new_paragraph.style = paragraph.style if paragraph.style is not None else "Normal"
    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    new_paragraph.paragraph_format.space_after = Pt(space_after)
    new_paragraph.paragraph_format.line_spacing = 1.25
    new_paragraph.paragraph_format.first_line_indent = Cm(0.74) if indent else Cm(0)
    run = new_paragraph.add_run(text)
    set_run_blue(run, bold=bold)
    return new_paragraph


def add_blue_blocks_after(paragraph, blocks):
    anchor = paragraph
    for block in blocks:
        if isinstance(block, tuple):
            text, bold, indent = block
        else:
            text, bold, indent = block, False, False
        anchor = add_blue_paragraph_after(anchor, text, bold=bold, indent=indent)
    return anchor


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def fill_blue_cell(cell, text, *, header=False):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(str(text))
    set_run_blue(run, bold=header, size=9.5)
    if header:
        shade_cell(cell, LIGHT_BLUE_FILL)


def add_blue_table_after(document, paragraph, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for index, header in enumerate(headers):
        fill_blue_cell(table.rows[0].cells[index], header, header=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            fill_blue_cell(cells[index], value)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    paragraph._p.addnext(table._tbl)
    return table


def find_first(document, predicate):
    for paragraph in document.paragraphs:
        if predicate(paragraph.text.strip()):
            return paragraph
    raise ValueError("Paragraph not found")


def add_note(document, predicate, blocks, table=None):
    paragraph = find_first(document, predicate)
    anchor = add_blue_blocks_after(paragraph, blocks)
    if table:
        add_blue_table_after(document, anchor, *table)


def add_note_after_table(document, table_index, blocks, table=None):
    table = document.tables[table_index]
    anchor = Paragraph(table._tbl, table._parent)
    # A table is not a paragraph. Insert a temporary paragraph after it and use it as the anchor.
    element = OxmlElement("w:p")
    table._tbl.addnext(element)
    anchor = Paragraph(element, table._parent)
    for block in blocks:
        if isinstance(block, tuple):
            text, bold, indent = block
        else:
            text, bold, indent = block, False, False
        run = anchor.add_run(text)
        set_run_blue(run, bold=bold)
        anchor.paragraph_format.space_after = Pt(3)
        anchor.paragraph_format.line_spacing = 1.25
        if blocks.index(block) != len(blocks) - 1:
            anchor = add_blue_paragraph_after(anchor, "")
    if table:
        add_blue_table_after(document, anchor, *table)


def main():
    document = Document(SOURCE)

    add_note(
        document,
        lambda text: text.startswith("���Ƿ��е���Ѳ����ô�ͳ��������߹ҹ�"),
        [
            ("�����ֲ��䣨���˶ԣ������鱣���ǰ���ƣ������Ӹ����⣺������������߳����ġ��̶���֪���ƶ�Ѳ�졪ͳһ�ܿء�ԭ��ϵͳ��", True, False),
            "����Ŀ��ǰ�ص��ǵ�·�����밲ȫѲ�죬��Ӧֱ�������Ѿ߱�������·/�����豸Ѳ�����������������Ƹ˵�ѹ����������ǻ�����״̬�ȴ�����������չΪ�����������С���·������ʩ����״̬��֪�����ӳ�����",
        ],
    )

    add_note(
        document,
        lambda text: text.startswith("������ص�������ʾĿǰ����Ŀ���о�״����"),
        [
            ("�����ֲ��䣨���˶ԣ�����ǰ�о��ձ齫��·������֪�����˻��Ӿ�Ѳ����ǻ�·�Ʒֱ���Ϊ��������չ��", True, False),
            "�̶����о��ص�ͨ�������󡢻�ˮ���˳��ӽ���������⣻���˻��о��ص�ͨ���Ǻ���ȡ֤����·����ʶ��͵Ϳ�Ѳ�죻�ǻ�·���о��ص�ͨ�����������ơ���������ͻ�����Ϣ���������з����Ĺ��Բ����ǹ̶������ݡ��ƶ�Ѳ�����ݺ͹������������ɢ�������γɿ�׷�ݵ�����ջ���",
        ],
        table=(
            ["�о�����", "�����о�״̬", "����Ŀ��Ӧ����"],
            [
                ["�̶��˻�����֪", "�ɼ����󡢻�ˮ������Ŀ���״̬�������澯", "ESP32-S3 ���� LD06��ˮλģ�顢RS485 ����վ�ͱ�����"],
                ["�Ϳյ�·Ѳ��", "ͨ�����˻����ջ���Ƶ�ɼ�·��ͼ��", "Hula ���˻�ִ����ɡ��ֶ�Ѳ������Ƶ�ش��������뷵��"],
                ["·���Ӿ�����", "���ô�ͳͼ��������ѧϰʶ���ѷ�/����", "���� OpenCV �������������ɱ�עͼ�� CSV ��¼"],
                ["�ۺϹܿ�ƽ̨", "���豸״̬����չʾ���ṩ�������", "Vue ���ͨ�� MQTT �� Flask �ӿ�ͳһչʾ�Ϳ���"],
            ],
            [3.2, 6.5, 7.2],
        ),
    )

    add_note(
        document,
        lambda text: text.startswith("���ⴴ�µ��Ƿ�ɲ���"),
        [
            ("�����ֲ��䣨���˶ԣ�������Ŀ����Χ�ơ�ϵͳ���ɷ���������ջ�����չ���£���������δ��ɼ���ǰֱ��д�ɡ������״�����Ψһ����", True, False),
            "������¹ؼ��ʣ��ǻ�·�ơ���Դ������֪�����˻�Ѳ�졢�Ӿ���־������·���ѷ�ʶ�𡢹̶������ƶ���Эͬ�����ص�˶Ե��������Ϊ���̶��˷��ո�֪���������˻��ֶ�Ѳ�����ա�������λ�롢���䶨λ�롢ͼ���ѷ�����������ҳ���������",
        ],
    )

    add_note(
        document,
        lambda text: text.startswith("���˻�������pyhulax SDK�����п��ơ�����滮��"),
        [
            ("�����ֲ��䣨���˶ԣ������˻�������ײ�ɿ����豸�����ṩ����Ŀ���������зɿء���Ŀ������ɵ��ǻ��� pyhulax SDK ������㿪����������״̬��ȡ���ֶκ�����֯�����/��ͣ/ǰ��/ת��/����ָ����š���־ʶ�𴥷�����Ƶ�����ս��롢�ѷ�������ã��Լ���ҳ������״̬չʾ��", True, False),
        ],
    )

    add_note(
        document,
        lambda text: text.startswith("��֪�㣨�̶���ʵ�ﹹ�ɣ�����ʵ�����װ�Ƿ�Ľ�"),
        [
            ("�����ֲ��䣨���˶ԣ�����ǰ�������ÿ���ʽ���ͲĽṹ�����ڴ��չʾ�����ߵ��Ժ�ģ���滻������ʵ�ʵ�·����ʱ�������Ϊ��·�Ƹ�/֧�� + ��ˮ����� + ������ƽ̨����һ�廯�ṹ��", True, False),
            "�̶��˿ɽ� ESP32-S3����Դת����MQTT ͨ������߶�����������ˮ�䣻LD06������վ�ͱ����Ʊ�����¶��ʹ��ר�÷����֣�ˮλ̽ͷ�����ݵ�ʵ�ʰ�װ���������ܱ���ģ�黯ά����Ҳ������԰����·������������ʩ�������Ӧ����ʽ��",
        ],
    )

    add_note(
        document,
        lambda text: text.startswith("������LED·��ģ�飨����+״ָ̬ʾ��"),
        [
            ("�����ֲ��䣨���˶ԣ������齫�������޶�Ϊ���������� LED ·��/״ָ̬ʾģ�飻������������Լ 40 cm �� 40 cm �����˻���ƽ̨��ƽ̨�ϲ��ý��䶨λ�룬��������Ŀ�ı߽��ʶ���ɲ�ж���ǺͰ�ȫ����������", True, False),
            "ԭ�ͽ׶��Ա���չʾ�͵���Ϊ������ʽӦ��ʱӦ���÷��������硢������ƽ̨���棬��ȷ����־��ߴ硢���������ͷ�ӳ���ƥ�䡣",
        ],
    )

    add_note(
        document,
        lambda text: text.startswith("RS485С�ͳ�����һ��ʽ����վ�ɼ�"),
        [
            ("�����ֲ��䣨���˶ԣ�����RS485 + Modbus-RTU���Ǵ���һ��ʽ����վ���õĽ��뷽ʽ�����б����ɱ��������ʽ���ǰӦ��ʵ���豸˵����˶Դ�վ��ַ�������ʡ�У��λ���Ĵ�����ַ�����١�", True, False),
            "Ŀǰϵͳ��Ԥ���ʵ������վ���ݲɼ�����ҳչʾ�ӿڡ�����վʵ������Ӧ������������¼��֤���١�������ʪ�ȡ���ѹ�����պ��������ֶ��Ƿ����豸�ֲ�һ�¡�",
        ],
    )

    add_note(
        document,
        lambda text: text.startswith("Hula���˻���������𣿣�"),
        [
            ("�����ֲ��䣨���˶ԣ�����ǰ����Ϊ�߾� Hula ϵ�п���/�������˻��������ͺ����Ի������ơ��ɹ���¼��ٷ��豸��ϢΪ׼���ĵ��н���ͻ������֤��������֧�� pyhulax SDK ���ӡ�ǰ������ͷ��Ƶ������ɡ���ͣ���ֶ�Ѳ�������ա���־ʶ���뽵�䡣", True, False),
        ],
    )

    add_note(
        document,
        lambda text: text.startswith("ͨ�Ų㣺��רҵ��̫ǿ"),
        [
            ("�����ֲ��䣨���ͨ�ױ�������ͨ�Ų�������Ϊϵͳ�ġ���Ϣͨ�������̶��˸���Ѵ���������������ҳ�����˻�������ձ��ؿ���ָ��ش���Ƶ/����״̬����ҳ����Щ��Ϣ������ʾ��ʹ���ߡ�", True, False),
            "���ʱ�ɲ�չ������Э��ϸ�ڣ�ֻ��˵����MQTT �ʺ��ö���̶��豸�����ϱ�С���ݣ�HTTP �ӿ��ʺ���ҳ�����ť�����˷���һ����ȷ������ָ����˻�����ͨ�������� TCP/SDK ��ɡ�",
        ],
    )

    add_note(
        document,
        lambda text: text.startswith("�̶���ͨ��MQTTЭ���ϴ�״̬���ݵ���ҳ��"),
        [
            ("�����ֲ��䣨���ͨ�ױ��������̶��˻ᶨʱ�ѡ�������롢��ˮ�������������ݡ�����״̬����С���ݷ��͵�ͳһ��Ϣͨ������ҳ���ĺ󼴿��Զ�ˢ�¡����������൱����ÿյ�����������������ƽ��/���쳣����", True, False),
        ],
    )

    add_note(
        document,
        lambda text: text.startswith("3.2.4 ��������"),
        [
            ("�����ֲ��䣨���˶ԣ������в���Ϊԭ����ƽ׶ε�ѡ����Ϣ����ʽ�汾Ӧ��ʵ�ʲɹ��ͺš���Ʒ˵���顢���ƺ��ֳ����Լ�¼Ϊ׼���ر�������վ���̡�LD06 ���ָ�ꡢ���˻��ͺ�/�������𽵱�־���ͺ�Ӧ���������ٶ��ⷢ����", True, False),
        ],
    )

    add_note(
        document,
        lambda text: text.startswith("3.3.2 �ɱ���Ԥ��"),
        [
            ("�����ֲ��䣨���˶ԣ������гɱ������Ϊ��ԭ�͹���Ԥ�㡱����Ӧֱ�ӵ�ͬ��ʵ�ʲɹ����㡣���齫�����豸�������豸�������ɹ��豸�ֿ�ͳ�ƣ��������Ʊ/������ͼ��Ϊ������", True, False),
        ],
        table=(
            ["Ԥ��ھ�", "���/�����ʽ", "�˶�Ҫ��"],
            [
                ["ԭ�ͺ������Ϲ���", "�����б�ϼ�Լ 5,460 Ԫ", "����˶��ͺš��������������Ƿ��ѹ�"],
                ["�����ظ��ɹ��豸", "���˻���ƽ�塢��ݮ�ɵȿɵ���Ϊ����/�����豸", "���⽫ѧУ/�ŶӼ����豸��дΪ����֧��"],
                ["Ԥ������", "���鰴�������Ϲ����Լ 10% Ԥ���Լ 546 Ԫ", "�������¡�ת�Ӱ塢��־�塢��ӡ���ṹ�������"],
                ["ԭ��Ԥ������", "Լ 6,006 Ԫ��5,460 + 546��", "�������Ԥ��˶ԣ����ս����ʵ�ʵ���Ϊ׼"],
            ],
            [3.5, 5.0, 8.4],
        ),
    )

    add_note(
        document,
        lambda text: text.startswith("�塢�������й���"),
        [
            ("�����ֲ��䣨���˶ԣ�������Ӧͻ������Ʒ��δӵ�ģ���ߵ�����ʾԭ�͡��ĵ������̣���������ͳ�������С����鰴���°汾·�����֡�", True, False),
        ],
        table=(
            ["�����׶�", "�������", "�γɵĿ���֤�ɹ�"],
            [
                ["V0 ������֤", "��ȷ�̶���֪�����˻�Ѳ�졢��ҳ�ܿ�����ܹ�", "���Ӳ���嵥��ͨ����·�볡���滮"],
                ["V1 �̶��˽���", "ESP32-S3 ���� LD06��ˮλ���������� MQTT", "��ҳ��ʾ��Ա�ӽ�����ˮ���豸״̬"],
                ["V2 ������֪�����", "Ԥ��/���� RS485 ����վ�����ƶ�����빦�ܿ��ؽ���", "��ɼ���չʾ�������·�·��"],
                ["V3 ���˻�����", "���� pyhulax SDK����Ƶ�������ա���־ʶ��ͽ���", "�����ɡ�Ѳ��������������ջ�����"],
                ["V4 �Ӿ�������������", "���պ�ִ���ѷ���������ɱ�עͼ�� CSV", "��ҳչʾ�����ѷ챨��������״̬"],
                ["V5 ��������", "Ǩ������ݮ�ɡ�·�����������ֳ������ṹ�Ż�", "����ʵ����ʾ��С��Χ�ֳ�����"],
            ],
            [3.0, 7.0, 6.9],
        ),
    )

    add_note(
        document,
        lambda text: text.startswith("5.2 ԭ�͵����׶�"),
        [
            ("�����ֲ��䣨���˶ԣ������齫���ڶ�λΪ������ԭ����֤��������д���Ѵ��ģ���̻����𡣵�ǰԭ������ɹ̶��˲ɼ�����ҳ���ơ����˻��Ӿ�����־�������ѷ�����Ĺ��ܱջ����ȶ��ԡ����ӹ�����Ӧ�ԡ����⿹����ͳ�����ͨ����������һ�׶��Ż����ݡ�", True, False),
        ],
    )

    add_note(
        document,
        lambda text: text.startswith("���貹�䡿 ���ṩ���׶εľ���ʱ����"),
        [
            ("�����ֲ��䣨���˶ԣ������Ȱ����׶Ρ�֤�ݡ����ݡ�������ϣ������������ں˶Ժ��滻Ϊ��ʵ���ڡ�", True, False),
            "����鵵����Ƭ/��Ƶ�������̶�������ṹ�������������ߡ���ҳ������塢���˻���Ƶ�������/Ѳ��/����/������̡��ѷ�ԭʼ��Ƭ���ע���ͼ��",
        ],
        table=(
            ["�׶�", "�������֤��", "�ؼ�����/�ļ�"],
            [
                ["�̶��˵���", "LD06��ˮλ�������ơ�ESP32 ��������ҳ��ͼ", "MQTT �ϱ���ͼ������/ADC/���յȼ���¼"],
                ["������֪", "����վ��װ����������ͼ", "RS485 �������Ĵ����������������¼"],
                ["���˻���������", "��ɡ���ͣ���ֶ�ǰ�������������Ƶ", "����״̬��������ToF �߶ȡ�������־"],
                ["�Ӿ��뵼��", "��Ƶ����2 �Ž������ 7 �ŷ�������Ƭ", "ʶ����������׶Ρ���Ƭ�浵"],
                ["�ѷ����", "ԭʼ��Ƭ����עͼ����ҳ�����ͼ", "media/photos��media/crack_results��result.csv"],
            ],
            [3.0, 7.0, 6.9],
        ),
    )

    add_note(
        document,
        lambda text: text.startswith("6.3 SWOT ����"),
        [
            ("�����ֲ��䣨���˶ԣ���SWOT �Ƿ����������ߣ���ǰ���ݿ����ڴ�磬����Ҫ�����Ԥ�⵱����ʵ�����齫������/���ơ���������������ݶ�Ӧ����������/��в��дΪ�ⲿ��������Ŀ���ա�", True, False),
            "���У��ѷ���׼ȷ�ʡ����˻�����ƫ��豸��������ʱ���ȶ���ָ�꣬Ӧ��ͳһ���Թ����²�������д�� SWOT �ġ�����֤�ݡ���",
        ],
    )

    add_note(
        document,
        lambda text: text.startswith("7.3 ���Ա���"),
        [
            ("�����ֲ��䣨���˶ԣ������Ա���Ӧ���֡���������ͨ���͡�����������������ǰ��ȷ�ϵ��Ƕ�ģ�鹦�ܱջ�����ɣ���Ҫ�����������ͳһ�����µ��ظ��������ݡ�ͳ�ƽ���ͻ����Աȡ�", True, False),
        ],
    )

    add_note(
        document,
        lambda text: text == "ϵͳ��������º��Ĺ��ܵĲ�����֤��",
        [
            ("�����ֲ��䣨���˶ԣ������������������֤��¼���¡����С�����֤����ʾ�������̿�ִ�У�����ͬ�ڶԾ��Ȼ򹤳̿ɿ��Ե����ճ�ŵ��", True, False),
        ],
        table=(
            ["ģ��", "����֤����", "��ǰ�����֤��", "���貹�����������"],
            [
                ["LD06 �˳��ӽ�", "���Ŀ����롢���յȼ���ƹ�״̬�������ʾ", "������/CRC ״̬��ͼ��MQTT �ϱ���¼", "��ͬ���롢��Ƕ�Ŀ���µ������©����"],
                ["��ˮģ��", "ADC ��������ʪ��ֵ��У���ˮ������ʾ", "��ʪ�궨ֵ����ҳ��ˮ״̬��ͼ", "ʵ��ˮ�����������Ķ�Ӧ����"],
                ["����վ", "����� RS485/Modbus �����������ҳ�ֶ�Ԥ��", "�豸˵���顢�������������Ƭ", "ʵ�����ߺ��������������ռ�¼"],
                ["���˻�����", "��ɡ�Ѳ������Ƶ�����ա�7 �ŷ����롢2 �Ž������뽵������", "����״̬��ͼ����Ƶ����������Ƭ����־", "��ͬ�߶�/����/��־�����µ�ʶ���������ƫ��"],
                ["�ѷ�ʶ��", "�Զ����ա�ͼ���������עͼ�� CSV ���", "ԭͼ����עͼ��result.csv����ҳ����", "���˹���ע��������׼ȷ�ʡ�����ʺ�©����"],
            ],
            [2.5, 4.2, 5.4, 4.8],
        ),
    )

    add_note(
        document,
        lambda text: text.startswith("���貹�䡿 ���鲹�䣺"),
        [
            ("�����ֲ��䣨���˶ԣ����������ͳһ���Ա�������ݣ�����ֻ��¼���ɹ�/ʧ�ܡ���ÿ�β���Ӧ��¼ʱ�䡢�ص㡢���ա��豸�汾��������ԭʼ���ݡ�������쳣ԭ�򣬲������Ӧ��ͼ/��Ƭ��", True, False),
            "��ǰ��������û���ظ��������˹���ע���յ��������д����ྫ�� ��X cm�����ѷ�ʶ��׼ȷ�� X%�������˻�����ƫ�� X cm���Ƚ��������֡�",
        ],
        table=(
            ["������Ŀ", "������������", "����ͳ��ָ��"],
            [
                ["LD06 ���", "3 ������� �� ÿ�� 10 ��", "ƽ�����������/©������"],
                ["��ˮ����", "��/ǳ/��/�� 4 ��ˮλ �� ÿ�� 5 ��", "ADC ��ֵ�������������ʵ��ˮ��ƫ��"],
                ["����ɼ�", "�������в����� 30 ����", "�ֶ������ʡ������ʡ���ο��豸ƫ��"],
                ["��־ʶ���뽵��", "ÿ�ָ߶�/����/�������� 10 ��", "ʶ��ɹ��ʡ���������ʡ����ƫ��"],
                ["�ѷ����", "���ѷ�/���ѷ�ͼƬ�������� 30 ��", "׼ȷ�ʡ�����ʡ�©���ʡ������ʱ"],
            ],
            [3.2, 5.6, 8.1],
        ),
    )

    # Table 7 is the SWOT table. Add a blue validation note without changing the original red content.
    add_note_after_table(
        document,
        6,
        [
            ("�����ֲ��䣨���˶ԣ���SWOT ���С����ᡱ�͡���в��������ҵ����������жϣ���������ʽ�����и��������ļ����г������Ʒ�������ϵ���Դ�����С����ơ��͡����ơ�Ӧ�뱾��Ŀ�Ĳ��Լ�¼�����й����嵥һһ��Ӧ��", True, False),
        ],
    )

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
